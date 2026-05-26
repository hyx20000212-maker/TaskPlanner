"""
Word Parser — Extracts text from .docx files using python-docx.
"""

from docx import Document as DocxDocument

from doc_parser.models import ParsedDocument


def parse_docx(file_path: str) -> ParsedDocument:
    """
    Parse a Word (.docx) file, extracting text paragraph by paragraph.

    Args:
        file_path: Path to the .docx file

    Returns:
        ParsedDocument with full text and paragraph/table counts

    Raises:
        ValueError: File is not a valid .docx format
    """
    try:
        doc = DocxDocument(file_path)
    except Exception as e:
        raise ValueError(f"Failed to open Word document: {e}") from e

    paragraphs: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                paragraphs.append(row_text)

    full_text = "\n\n".join(paragraphs)

    if not full_text.strip():
        raise ValueError("No text extracted from the Word document.")

    return ParsedDocument(
        source=file_path,
        file_type="docx",
        raw_text=full_text,
        page_count=None,  # python-docx does not expose page count
        metadata={
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
        },
    )
